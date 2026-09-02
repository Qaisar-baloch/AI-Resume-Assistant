import io
import json
import os
import re
from typing import Any, Dict, List

import streamlit as st
from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document


APP_TITLE = "AI Resume ATS Checker"
DEFAULT_MODEL = "gemini-3.7-flash"
MAX_RESUME_CHARS = 50000


def get_api_key() -> str | None:
    """Read the Gemini API key from Streamlit secrets or an environment variable."""
    try:
        key = st.secrets.get("GEMINI_API_KEY")
        if key:
            return str(key)
    except Exception:
        pass

    key = os.getenv("GEMINI_API_KEY")
    return key.strip() if key else None


def extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages).strip()


def extract_docx_text(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    parts: List[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_resume_text(uploaded_file) -> str:
    data = uploaded_file.getvalue()
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        return extract_pdf_text(data)
    if name.endswith(".docx"):
        return extract_docx_text(data)
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="replace").strip()

    raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")


def clean_json_text(text: str) -> str:
    """Remove accidental Markdown fences if the model returns them."""
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)
    return text.strip()


def analyze_resume(resume_text: str, model_name: str) -> Dict[str, Any]:
    client = genai.Client(api_key=get_api_key())

    prompt = f"""
You are an expert ATS (Applicant Tracking System) resume reviewer and career coach.

Analyze the resume below as a GENERAL-PURPOSE ATS resume, not against a specific job description.

Important:
- Do not invent experience, education, skills, employers, dates, metrics, certifications, or achievements.
- Treat an ATS score as an estimate, not a real score from a particular employer's ATS.
- Reward clear standard headings, searchable keywords, measurable achievements, relevant skills,
  concise bullets, consistent dates, and ATS-friendly formatting.
- Penalize missing/unclear contact information, vague bullets, excessive length, poor structure,
  tables/text boxes/graphics if they are evident from the extracted text, keyword stuffing, and
  grammar/clarity issues.
- Give practical improvements the applicant can actually make.

Return ONLY valid JSON with exactly this structure:
{{
  "ats_score": 0,
  "score_label": "Poor|Needs Improvement|Good|Very Good|Excellent",
  "summary": "short overall assessment",
  "strengths": ["...", "...", "..."],
  "improvements": [
    {{
      "priority": "High|Medium|Low",
      "issue": "...",
      "recommendation": "...",
      "example": "..."
    }}
  ],
  "ats_keywords": ["keyword", "..."],
  "formatting_checks": [
    {{
      "check": "...",
      "status": "Pass|Needs Attention",
      "comment": "..."
    }}
  ]
}}

Scoring guidance:
- 90-100: Excellent ATS readiness
- 80-89: Very Good
- 70-79: Good
- 60-69: Needs Improvement
- 0-59: Poor

Resume:
---BEGIN RESUME---
{resume_text[:MAX_RESUME_CHARS]}
---END RESUME---
"""

    response = client.models.generate_content(
        model=model_name,
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            temperature=0.2,
        ),
    )

    raw = clean_json_text(response.text or "")
    result = json.loads(raw)

    # Defensive validation so the UI never trusts an invalid model response.
    score = int(result.get("ats_score", 0))
    result["ats_score"] = max(0, min(100, score))

    if not isinstance(result.get("strengths"), list):
        result["strengths"] = []
    if not isinstance(result.get("improvements"), list):
        result["improvements"] = []
    if not isinstance(result.get("ats_keywords"), list):
        result["ats_keywords"] = []
    if not isinstance(result.get("formatting_checks"), list):
        result["formatting_checks"] = []

    return result


def main() -> None:
    st.set_page_config(
        page_title=APP_TITLE,
        page_icon="📄",
        layout="wide",
    )

    st.title("📄 AI Resume ATS Checker")
    st.caption(
        "Upload a resume and Gemini Flash will estimate ATS readiness and suggest "
        "specific improvements."
    )

    with st.sidebar:
        st.header("Settings")
        model_name = st.text_input(
            "Gemini model",
            value=os.getenv("GEMINI_MODEL", DEFAULT_MODEL),
            help="You can change this if Google changes model availability.",
        )
        st.divider()
        st.markdown(
            "**Supported files:** PDF, DOCX, TXT\n\n"
            "**Privacy:** Your resume is sent to Gemini for analysis. "
            "Do not upload sensitive information you do not want processed by an external AI service."
        )

    uploaded_file = st.file_uploader(
        "Upload your resume",
        type=["pdf", "docx", "txt"],
        help="For best results, use a text-based PDF or DOCX.",
    )

    if not uploaded_file:
        st.info("Upload a resume to start the ATS analysis.")
        return

    st.success(f"Loaded: {uploaded_file.name}")

    if st.button("🔍 Analyze Resume", type="primary", use_container_width=True):
        api_key = get_api_key()
        if not api_key:
            st.error(
                "Gemini API key not found. Add GEMINI_API_KEY to Streamlit Secrets "
                "or set it as an environment variable."
            )
            st.stop()

        try:
            with st.spinner("Reading and analyzing your resume..."):
                resume_text = extract_resume_text(uploaded_file)

                if not resume_text:
                    st.error(
                        "No readable text was found. If this is a scanned/image-only PDF, "
                        "use an OCR-enabled PDF or upload a DOCX/TXT version."
                    )
                    st.stop()

                result = analyze_resume(resume_text, model_name)
                st.session_state["analysis"] = result

        except json.JSONDecodeError:
            st.error("Gemini returned an unexpected response format. Please try again.")
            st.stop()
        except Exception as exc:
            st.error(f"Analysis failed: {exc}")
            st.stop()

    result = st.session_state.get("analysis")
    if not result:
        return

    st.divider()

    col1, col2 = st.columns([1, 2])
    with col1:
        st.metric("Estimated ATS Score", f"{result['ats_score']}/100")
    with col2:
        st.subheader(result.get("score_label", "ATS Assessment"))
        st.write(result.get("summary", ""))

    st.subheader("✅ Strengths")
    strengths = result.get("strengths", [])
    if strengths:
        for item in strengths:
            st.markdown(f"- {item}")
    else:
        st.write("No strengths were returned.")

    st.subheader("🛠️ Improvements")
    improvements = result.get("improvements", [])
    if improvements:
        for item in improvements:
            priority = item.get("priority", "Medium")
            issue = item.get("issue", "")
            recommendation = item.get("recommendation", "")
            example = item.get("example", "")

            with st.container(border=True):
                st.markdown(f"**{priority} priority — {issue}**")
                st.write(recommendation)
                if example:
                    st.caption(f"Example: {example}")
    else:
        st.write("No improvement suggestions were returned.")

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("🔑 ATS Keywords")
        keywords = result.get("ats_keywords", [])
        st.write(", ".join(map(str, keywords)) if keywords else "None identified.")

    with col2:
        st.subheader("📋 Formatting Checks")
        checks = result.get("formatting_checks", [])
        for check in checks:
            status = check.get("status", "Needs Attention")
            icon = "✅" if status == "Pass" else "⚠️"
            st.markdown(
                f"{icon} **{check.get('check', '')}** — {check.get('comment', '')}"
            )

    st.caption(
        "Note: This score is an AI-generated estimate. Different employers and ATS platforms "
        "use different parsing and ranking methods."
    )


if __name__ == "__main__":
    main()
